"""DOCX and PDF, both walking the shared block model.

Single column, no tables, no text boxes, no graphics, standard headings,
selectable text. Everything an ATS can actually read — the formats that look
sophisticated are the ones that parse to nothing.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from .layout import MARGIN_INCHES, TYPE, build_blocks, line_height_for, spacing_for


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def write_docx(doc: dict[str, Any], path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt

    blocks = build_blocks(doc)
    space = spacing_for(blocks)
    leading = line_height_for(blocks)

    document = Document()
    section = document.sections[0]
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(MARGIN_INCHES))

    normal = document.styles["Normal"]
    # Arial pairs with Helvetica in the PDF: metrically equivalent, so the two
    # documents look the same without either shipping a font file.
    normal.font.name = "Arial"
    normal.font.size = Pt(TYPE["body"])

    usable = section.page_width - section.left_margin - section.right_margin

    def para(after: float, before: float = 0.0):
        p = document.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_after = Pt(after)
        fmt.space_before = Pt(before)
        fmt.line_spacing = leading
        return p

    for block in blocks:
        if block.kind == "name":
            run = para(space["after_name"]).add_run(block.text)
            run.bold = True
            run.font.size = Pt(TYPE["name"])
        elif block.kind == "headline":
            para(space["after_headline"]).add_run(block.text).font.size = Pt(TYPE["headline"])
        elif block.kind == "contact":
            para(space["after_contact"]).add_run(block.text).font.size = Pt(TYPE["contact"])
        elif block.kind == "section":
            run = para(space["after_section"], space["before_section"]).add_run(block.text)
            run.bold = True
            run.font.size = Pt(TYPE["section"])
        elif block.kind == "paragraph":
            para(space["between_paragraphs"]).add_run(block.text)
        elif block.kind == "skills":
            p = para(space["between_skill_lines"])
            p.add_run(block.label + ": ").bold = True
            p.add_run(block.text)
        elif block.kind == "role_header":
            p = para(space["after_role_header"], space["before_role"])
            # A right tab stop rather than a table: a table is the single most
            # common way a resume parses into nonsense.
            p.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
            p.add_run(block.left).bold = True
            if block.right:
                p.add_run("\t" + block.right)
        elif block.kind == "role_meta":
            run = para(space["after_role_meta"]).add_run(block.text)
            run.italic = True
            run.font.size = Pt(TYPE["meta"])
        elif block.kind == "bullet":
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(space["between_bullets"])
            p.paragraph_format.line_spacing = leading
            p.add_run(block.text)
        elif block.kind == "labelled":
            p = para(space["between_paragraphs"])
            p.add_run(block.label + ": ").bold = True
            p.add_run(block.text)

    document.save(str(path))


def write_pdf(doc: dict[str, Any], path: Path) -> None:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    blocks = build_blocks(doc)
    space = spacing_for(blocks)
    leading_ratio = line_height_for(blocks)

    def style(name: str, size: float) -> ParagraphStyle:
        return ParagraphStyle(
            name,
            # Helvetica is built into the PDF format, so there is no font to
            # fetch. Renderers tend to fail silently when a registered font
            # cannot be loaded, which is a miserable bug to chase.
            fontName="Helvetica",
            fontSize=size,
            leading=size * leading_ratio,
            alignment=TA_LEFT,
        )

    template = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        topMargin=MARGIN_INCHES * inch,
        bottomMargin=MARGIN_INCHES * inch,
        leftMargin=MARGIN_INCHES * inch,
        rightMargin=MARGIN_INCHES * inch,
        title=(doc.get("contact") or {}).get("fullName", "Resume"),
    )

    body = style("body", TYPE["body"])
    story: list[Any] = []
    pending: list[Any] = []

    def flush() -> None:
        if not pending:
            return
        story.append(
            ListFlowable(
                list(pending),
                bulletType="bullet",
                start="•",
                leftIndent=12,
                bulletFontSize=TYPE["body"],
                spaceAfter=0,
            )
        )
        pending.clear()

    for block in blocks:
        if block.kind != "bullet":
            flush()

        if block.kind == "name":
            story.append(Paragraph("<b>" + _escape(block.text) + "</b>", style("name", TYPE["name"])))
            story.append(Spacer(1, space["after_name"]))
        elif block.kind == "headline":
            story.append(Paragraph(_escape(block.text), style("headline", TYPE["headline"])))
            story.append(Spacer(1, space["after_headline"]))
        elif block.kind == "contact":
            story.append(Paragraph(_escape(block.text), style("contact", TYPE["contact"])))
            story.append(Spacer(1, space["after_contact"]))
        elif block.kind == "section":
            story.append(Spacer(1, space["before_section"]))
            story.append(
                Paragraph("<b>" + _escape(block.text) + "</b>", style("section", TYPE["section"]))
            )
            story.append(Spacer(1, space["after_section"]))
        elif block.kind == "paragraph":
            story.append(Paragraph(_escape(block.text), body))
            story.append(Spacer(1, space["between_paragraphs"]))
        elif block.kind == "skills":
            story.append(
                Paragraph("<b>" + _escape(block.label) + ":</b> " + _escape(block.text), body)
            )
            story.append(Spacer(1, space["between_skill_lines"]))
        elif block.kind == "role_header":
            story.append(Spacer(1, space["before_role"]))
            header = style("role", TYPE["body"])
            # Right-aligned dates via a tab stop, matching the DOCX. No table,
            # for the same reason as the DOCX.
            header.tabs = [(template.width, "RIGHT")]
            markup = "<b>" + _escape(block.left) + "</b>"
            if block.right:
                markup += "\t" + _escape(block.right)
            story.append(Paragraph(markup, header))
            story.append(Spacer(1, space["after_role_header"]))
        elif block.kind == "role_meta":
            story.append(
                Paragraph("<i>" + _escape(block.text) + "</i>", style("meta", TYPE["meta"]))
            )
            story.append(Spacer(1, space["after_role_meta"]))
        elif block.kind == "bullet":
            pending.append(
                ListItem(
                    Paragraph(_escape(block.text), body),
                    spaceAfter=space["between_bullets"],
                )
            )
        elif block.kind == "labelled":
            story.append(
                Paragraph("<b>" + _escape(block.label) + ":</b> " + _escape(block.text), body)
            )
            story.append(Spacer(1, space["between_paragraphs"]))

    flush()
    template.build(story)
